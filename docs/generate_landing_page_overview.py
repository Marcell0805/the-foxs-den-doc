"""Generate The Fox's Den landing page overview Word document."""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise SystemExit("Run: pip install python-docx")

OUT = Path(__file__).resolve().parent / "The Fox's Den - Landing Page Overview.docx"

doc = Document()

title = doc.add_heading("The Fox's Den", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Personal portfolio landing page — overview")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True

doc.add_paragraph()

# 1. Executive summary
doc.add_heading("1. Executive summary", level=1)
doc.add_paragraph(
    "The Fox's Den is a password-gated static portfolio site hosted on GitHub Pages. "
    "It serves as a single hub where visitors can browse personal websites and mobile apps, "
    "read short descriptions, download APKs or desktop tool packages, and open linked "
    "external sites. The landing page is the home screen: brand, search, grouped project "
    "lists, and contact links — all driven from JSON data and rebuilt by a PowerShell build script."
)

# 2. Why it was created
doc.add_heading("2. Why it was created", level=1)
doc.add_paragraph(
    "Several personal projects were growing in parallel — Flutter mobile apps (e.g. Active Huntress, "
    "Huntress Cookbook Mobile), a recipe website, and desktop utilities. Each needed a public "
    "place to host APKs, version manifests for in-app updates, and a professional entry point "
    "without maintaining a separate site per app."
)
doc.add_paragraph(
    "The Fox's Den was created to:"
)
for item in [
    "Centralise project discovery in one branded landing page",
    "Host APK and zip downloads on GitHub Pages (same pattern as The Huntress Cookbook hub)",
    "Publish mobile-version.json files so apps can check for updates in-app",
    "Separate Live and Beta download channels where needed",
    "Present websites and tools alongside mobile apps in a consistent layout",
    "Keep the site static (no server) — JSON + HTML + client-side search only",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "The design follows a cookbook-style portal model: manifest-driven listings, README-sourced "
    "detail pages, and a build step that regenerates navigation and search from data files."
)

# 3. What the landing page does
doc.add_heading("3. What the landing page does", level=1)
doc.add_paragraph(
    "When a visitor opens the site and enters the portal password, they land on index.html. "
    "The page shows:"
)
for item in [
    "Logo and title — The Fox's Den",
    "Tagline — e.g. “Websites and mobile apps to browse, install, and try.”",
    "Search field — opens a modal (Ctrl+K) with Fuse.js full-text search across projects",
    "Grouped listings — Mobile apps, Websites, and Tools, each numbered from 1 within its group",
    "“Show more” — lists beyond 10 items per group collapse behind an expand control",
    "Learn more — link to the About section",
    "Footer — email, GitHub, and LinkedIn from portal-settings.json",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "Clicking a listed project opens a detail page (sections/<id>.html) with status badge, "
    "summary, download buttons (APK or zip with version and file size), and About content "
    "from the project README or a curated summaryOverride."
)

# 4. Functionality
doc.add_heading("4. Functionality", level=1)

doc.add_heading("4.1 Landing page layout", level=2)
for item in [
    "Woodland-themed visual design (green/gold palette, Fraunces + Source Sans 3)",
    "Card-style groups for Mobile apps, Websites, and Tools with clear section headers",
    "Search bar aligned to the same width as listing cards",
    "Responsive layout — mobile drawer sidebar on detail pages",
    "Password gate on first visit (session storage; client-side only)",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4.2 Project detail pages", level=2)
for item in [
    "Sidebar navigation grouped by kind (Mobile / Websites / Tools); About pinned at the bottom",
    "Status badges — Live, Beta, In Progress, or Planned",
    "Mobile — Download Live APK and optional Beta APK with version, build, and file size",
    "Websites — Open site button linking to external GitHub Pages URL",
    "Tools — Download zip with version, size, and extract-and-run instructions",
    "About block — README excerpt or summaryOverride; markdown lightly cleaned for display",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4.3 Search", level=2)
doc.add_paragraph(
    "Ctrl+K (or clicking the landing search field) opens a modal. Fuse.js searches a local "
    "index built from project titles, summaries, tags, and README text. Results link directly "
    "to section pages. Hidden pages are excluded from the index by design."
)

doc.add_heading("4.4 About page", level=2)
for item in [
    "Short personal blurb and skills list (from portal-settings.json)",
    "Contact cards — email, GitHub, LinkedIn",
    "Included in sidebar nav and search like other sections",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4.5 Data and build pipeline", level=2)
doc.add_paragraph(
    "Content is not edited in HTML by hand for each app. Maintainers update "
    "portal/data/apps-manifest.json (and optionally project READMEs or summaryOverride fields), "
    "then run build-portal.ps1. The build script:"
)
for item in [
    "Syncs each app into portal/data/<id>.json and nav.json",
    "Generates portal/js/portal-data.js and search-index.js",
    "Writes section HTML shells under portal/sections/",
    "Reads APK/zip metadata from portal/downloads/ for version and size on detail pages",
    "Applies cache-bust query strings on script includes",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "GitHub Actions deploys the portal/ folder to GitHub Pages on push to main."
)

# 5. Listing types
doc.add_heading("5. Listing types", level=1)
table = doc.add_table(rows=4, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Kind"
hdr[1].text = "Landing group"
hdr[2].text = "Typical action on detail page"
rows = [
    ("mobile", "Mobile apps", "Download Live/Beta APK; in-app update via mobile-version.json"),
    ("website", "Websites", "Open external site (e.g. Huntress Cookbook on GitHub Pages)"),
    ("tool", "Tools", "Download self-contained Windows zip package"),
]
for i, (a, b, c) in enumerate(rows, start=1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b
    table.rows[i].cells[2].text = c

doc.add_paragraph()

# 6. Configuration
doc.add_heading("6. Configuration", level=1)
for item in [
    "portal/data/portal-settings.json — site name, tagline, contact, About blurb/skills, theme colours, pagesBaseUrl, password",
    "portal/data/apps-manifest.json — all project listings (id, title, kind, status, paths, APK/zip names, external URLs)",
    "portal/data/my-huntress.json — hidden dedication page (not in nav or search; separate unlock flow)",
    "portal/downloads/ — APKs, zips, and per-app version JSON files served publicly",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "Live URL (when deployed): https://marcell0805.github.io/the-foxs-den-doc/"
)

# 7. Out of scope / notes
doc.add_heading("7. Notes", level=1)
for item in [
    "Password protection is client-side only — suitable for a private portfolio, not strong security",
    "APKs and version JSON URLs are public once published",
    "Fox Publish (separate WinForms tool) automates deploys to this hub but is not part of the landing page itself",
    "AppGen can round-trip mobile app config with the portal folder for generated Flutter projects",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph()
p = doc.add_paragraph("Document describes The Fox's Den landing page and portal hub.")
p.runs[0].font.size = Pt(9)
p.runs[0].italic = True

doc.save(OUT)
print(f"Wrote {OUT}")
